"""
Генерація Word документів на основі шаблонів
"""
from docx import Document
from datetime import datetime
import os
import re
from typing import Dict, List, Tuple


class ReportGenerator:
    """
    Клас для генерації Word документів
    """

    def __init__(self, template_path: str):
        """
        Ініціалізація генератора

        Args:
            template_path: Шлях до шаблону Word
        """
        self.template_path = template_path

        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Шаблон не знайдено: {template_path}")

    def find_manual_markers(self, doc: Document) -> List[str]:
        """
        Знаходить всі MANUAL маркери в документі

        Args:
            doc: Об'єкт Document

        Returns:
            Список унікальних MANUAL маркерів
        """
        markers = set()
        pattern = r'\{\{MANUAL:([^}]+)\}\}'

        # Шукаємо в параграфах
        for paragraph in doc.paragraphs:
            if paragraph.text:
                found = re.findall(pattern, paragraph.text)
                markers.update(found)

        # Шукаємо в таблицях
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text:
                            found = re.findall(pattern, paragraph.text)
                            markers.update(found)

        return list(markers)

    def replace_placeholders(self, doc: Document, data: Dict, manual_data: Dict = None) -> None:
        """
        Заміняє плейсхолдери в документі на реальні дані

        Args:
            doc: Об'єкт Document
            data: Словник з даними для заміни
            manual_data: Словник з вручну введеними даними (для MANUAL маркерів)
        """
        if manual_data is None:
            manual_data = {}

        # Форматуємо ПІБ: прізвище великими літерами
        full_name = data.get("name", "")
        if full_name:
            name_parts = full_name.split()
            if name_parts:
                # Прізвище (перше слово) - великими літерами
                name_parts[0] = name_parts[0].upper()
                full_name = " ".join(name_parts)

        # Мапінг плейсхолдерів до даних
        placeholders = {
            "{{ПІБ}}": full_name,
            "{{ЗВАННЯ}}": data.get("rank", ""),
            "{{ПОСАДА}}": data.get("position", ""),
            "{{РНОКПП}}": data.get("rnokpp", ""),
            "{{Дата народження}}": data.get("birth_date", ""),
            "{{ПЕРІОДИ}}": data.get("periods", ""),
            "{{ПЕРІОДИ_100}}": data.get("periods_100", ""),
            "{{ПЕРІОДИ_30}}": data.get("periods_30", ""),
            "{{ПЕРІОДИ_ВСІ}}": data.get("periods_all", ""),
            "{{ЖБД}}": data.get("zbd", ""),
            "{{ГРОМАДА}}": data.get("hromady", ""),
            "{{ДАТА}}": datetime.now().strftime("%d.%m.%Y")
        }

        # Додаємо MANUAL маркери
        # Спеціальна обробка для серії та номера паспорту
        seria_empty = not manual_data.get("СЕРІЯ", "").strip()
        nomer_empty = not manual_data.get("НОМЕР", "").strip()

        for key, value in manual_data.items():
            # Якщо обидва поля паспорту порожні
            if key == "СЕРІЯ" and seria_empty and nomer_empty:
                placeholders[f"{{{{MANUAL:{key}}}}}"] = ""
            elif key == "НОМЕР" and seria_empty and nomer_empty:
                placeholders[f"{{{{MANUAL:{key}}}}}"] = "відсутній"
            else:
                placeholders[f"{{{{MANUAL:{key}}}}}"] = value

        # Заміна в параграфах
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, placeholders)

        # Заміна в таблицях
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, placeholders)

        # Після заміни всіх плейсхолдерів виділяємо періоди, ЖБД та громади
        self._format_special_content(doc, data.get("zbd", ""), data.get("hromady", ""))

    def _replace_in_paragraph(self, paragraph, placeholders):
        """
        Заміняє плейсхолдери в параграфі, навіть якщо вони розбиті між runs
        """
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                # Збираємо весь текст параграфу
                full_text = paragraph.text

                # Заміняємо плейсхолдер
                new_text = full_text.replace(placeholder, str(value))

                # Якщо текст змінився
                if new_text != full_text:
                    # Очищуємо подвійні пробіли
                    new_text = re.sub(r'\s{2,}', ' ', new_text)

                    # Зберігаємо форматування першого run
                    if paragraph.runs:
                        first_run = paragraph.runs[0]
                        # Очищаємо всі runs
                        for run in paragraph.runs:
                            run.text = ""
                        # Записуємо новий текст у перший run
                        first_run.text = new_text
                    else:
                        # Якщо немає runs, створюємо новий
                        paragraph.text = new_text

    def _format_special_content(self, doc: Document, zbd_text: str, hromady_text: str) -> None:
        """
        Форматує спеціальний контент: періоди, ЖБД, громади
        Обробляє все за один прохід, щоб уникнути конфліктів

        Args:
            doc: Об'єкт Document
            zbd_text: Текст ЖБД (підкреслення)
            hromady_text: Текст громад (жирний)
        """
        # Патерн для пошуку періодів
        period_pattern = r'з\s+\d{2}\.\d{2}\.\d{4}\s+по\s+\d{2}\.\d{2}\.\d{4}'

        # Обробляємо параграфи
        for paragraph in doc.paragraphs:
            self._format_paragraph_content(paragraph, period_pattern, zbd_text, hromady_text)

        # Обробляємо таблиці
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._format_paragraph_content(paragraph, period_pattern, zbd_text, hromady_text)

    def _format_paragraph_content(self, paragraph, period_pattern: str, zbd_text: str, hromady_text: str) -> None:
        """
        Форматує контент в одному параграфі

        Args:
            paragraph: Параграф для обробки
            period_pattern: Регулярний вираз для періодів
            zbd_text: Текст ЖБД (підкреслення)
            hromady_text: Текст громад (жирний)
        """
        text = paragraph.text

        # Знаходимо всі спеціальні елементи в тексті
        formatting_ranges = []

        # 1. Знаходимо періоди (жирний + підкреслення)
        period_matches = list(re.finditer(period_pattern, text))
        if period_matches:
            # Об'єднуємо всі періоди в один діапазон (від першого до останнього)
            first_start = period_matches[0].start()
            last_end = period_matches[-1].end()
            formatting_ranges.append((first_start, last_end, True, True))  # (start, end, bold, underline)

        # 2. ЖБД - без форматування (раніше було підкреслення)
        # ЖБД тепер виводиться без спеціального форматування

        # 3. Знаходимо громади (тільки жирний)
        if hromady_text and hromady_text.strip():
            hromady_pos = text.find(hromady_text)
            if hromady_pos != -1:
                formatting_ranges.append((hromady_pos, hromady_pos + len(hromady_text), True, False))

        # Якщо нічого не знайдено - виходимо
        if not formatting_ranges:
            return

        # Сортуємо діапазони за позицією початку
        formatting_ranges.sort(key=lambda x: x[0])

        # Зберігаємо оригінальне форматування
        original_font = None
        if paragraph.runs:
            first_run = paragraph.runs[0]
            if hasattr(first_run.font, 'name'):
                original_font = (first_run.font.name, first_run.font.size)

        # Очищаємо параграф
        paragraph.clear()

        # Формуємо новий параграф з форматуванням
        current_pos = 0

        for start, end, bold, underline in formatting_ranges:
            # Додаємо текст до форматованого блоку
            if start > current_pos:
                run = paragraph.add_run(text[current_pos:start])
                if original_font:
                    run.font.name = original_font[0]
                    run.font.size = original_font[1]

            # Додаємо форматований блок
            run = paragraph.add_run(text[start:end])
            if original_font:
                run.font.name = original_font[0]
                run.font.size = original_font[1]
            if bold:
                run.bold = True
            if underline:
                run.underline = True

            current_pos = end

        # Додаємо залишок тексту
        if current_pos < len(text):
            run = paragraph.add_run(text[current_pos:])
            if original_font:
                run.font.name = original_font[0]
                run.font.size = original_font[1]

    def generate_report(self, servicemember_data: Dict, output_path: str, manual_data: Dict = None) -> bool:
        """
        Генерує рапорт для одного військовослужбовця

        Args:
            servicemember_data: Словник з даними військовослужбовця
            output_path: Шлях для збереження згенерованого документу
            manual_data: Словник з вручну введеними даними

        Returns:
            True якщо успішно

        Raises:
            Exception: Будь-яка помилка при генерації (шаблон, збереження, тощо)
        """
        # Завантажуємо шаблон
        doc = Document(self.template_path)

        # Заміняємо плейсхолдери
        self.replace_placeholders(doc, servicemember_data, manual_data)

        # Зберігаємо документ
        doc.save(output_path)

        return True

    def batch_generate(
        self,
        data_list: List[Dict],
        output_dir: str,
        file_prefix: str = "Рапорт"
    ) -> Tuple[int, int]:
        """
        Генерує рапорти для пакета військовослужбовців

        Args:
            data_list: Список словників з даними
            output_dir: Директорія для збереження документів
            file_prefix: Префікс назви файлу

        Returns:
            Кортеж (успішно згенеровано, помилок)
        """
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        success_count = 0
        error_count = 0

        for data in data_list:
            # Генеруємо унікальне ім'я файлу
            name = data.get("name", "Невідомо")
            # Замінюємо небезпечні символи в імені файлу
            safe_name = name.replace(" ", "_").replace("/", "_").replace("\\", "_")

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{file_prefix}_{safe_name}_{timestamp}.docx"
            output_path = os.path.join(output_dir, filename)

            # Генеруємо рапорт
            if self.generate_report(data, output_path):
                success_count += 1
            else:
                error_count += 1

        return (success_count, error_count)
